# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\QLDAPM\Management-Application-Project-Group-2")
OUT = ROOT / "reports" / "Bao_cao_du_an_quan_ly_quan_ca_phe_incremental_fixed.docx"


def set_run_font(run, size=13, bold=False, italic=False, color=None, font="Times New Roman"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    set_run_font(run, size=12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbers(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_code(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10, font="Consolas")


def setup_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2)

    for style_name in ["Normal", "Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.styles["Normal"].font.size = Pt(13)
    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(13)
    doc.styles["Heading 3"].font.bold = True
    return doc


doc = setup_document()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("BÁO CÁO DỰ ÁN PHẦN MỀM")
set_run_font(r, size=18, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("HỆ THỐNG QUẢN LÝ QUÁN CÀ PHÊ - ARTISANAL CAFE")
set_run_font(r, size=18, bold=True, color=RGBColor(47, 84, 150))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Áp dụng mô hình phát triển Incremental")
set_run_font(r, size=14, italic=True)

doc.add_paragraph("")
add_table(
    doc,
    ["Thông tin", "Nội dung"],
    [
        ("Dự án", "Management Application Project - Group 2"),
        ("Phạm vi báo cáo", "Version 1: Quản lý Menu; Version 2: Tạo Order cốt lõi"),
        ("Công nghệ", "Java Spring Boot, Thymeleaf, Spring Data JPA, MySQL, HTML/CSS/JavaScript"),
        ("Người thực hiện", "Sinh viên phụ trách Version 1 và Version 2"),
        ("Ngày lập báo cáo", "19/05/2026"),
    ],
)

doc.add_page_break()
doc.add_heading("MỤC LỤC TÓM TẮT", level=1)
add_numbers(
    doc,
    [
        "Giới thiệu dự án",
        "Công nghệ sử dụng và cấu trúc mã nguồn",
        "Áp dụng mô hình Incremental",
        "Product Backlog",
        "Definition of Done",
        "Version 1: Quản lý Menu",
        "Version 2: Tạo Order cốt lõi",
        "Thiết kế cơ sở dữ liệu và kiến trúc xử lý",
        "Kiểm thử, đánh giá hiện trạng và hướng phát triển",
        "Kết luận",
    ],
)

doc.add_page_break()
doc.add_heading("1. Giới thiệu dự án", level=1)
doc.add_heading("1.1. Tên dự án", level=2)
doc.add_paragraph("Hệ thống quản lý quán cà phê Artisanal Cafe - Coffee Management Application.")

doc.add_heading("1.2. Mục tiêu dự án", level=2)
doc.add_paragraph(
    "Dự án xây dựng một ứng dụng web hỗ trợ vận hành cơ bản cho quán cà phê. "
    "Ứng dụng tập trung vào hai nhóm nghiệp vụ quan trọng ở giai đoạn đầu: quản lý thực đơn và bán hàng tại quầy POS. "
    "Sau mỗi phiên bản, hệ thống phải tạo ra một increment có thể chạy được, có thể kiểm tra bằng trình duyệt và dữ liệu được lưu vào cơ sở dữ liệu."
)
add_bullets(
    doc,
    [
        "Quản trị viên có thể quản lý danh sách món uống/món ăn trong thực đơn.",
        "Sản phẩm được phân theo danh mục để dễ lọc và chọn món.",
        "Nhân viên có thể mở giao diện POS, chọn món, lập đơn và thanh toán tại quầy.",
        "Đơn hàng sau khi thanh toán được lưu xuống MySQL để phục vụ lịch sử đơn hàng và báo cáo.",
        "Hệ thống có nền tảng để mở rộng sang thanh toán điện tử, báo cáo doanh thu và quản lý hóa đơn.",
    ],
)

doc.add_heading("1.3. Phạm vi cá nhân được giao", level=2)
add_table(
    doc,
    ["Version", "Tên version", "Nội dung phụ trách", "Giá trị sau khi hoàn thành"],
    [
        (
            "Version 1",
            "Quản lý Menu",
            "Setup project. Hoàn thành chức năng thêm/sửa/xóa sản phẩm và tích hợp danh mục vào thực đơn.",
            "Quản trị viên có thể quản lý menu để chuẩn bị dữ liệu bán hàng.",
        ),
        (
            "Version 2",
            "Tạo Order cốt lõi",
            "Xây dựng giao diện POS cho nhân viên, chọn món, lập giỏ hàng và checkout tiền mặt.",
            "Quán có thể dùng thử để tính tiền và lưu đơn hàng cơ bản.",
        ),
    ],
)

doc.add_heading("2. Công nghệ sử dụng và cấu trúc mã nguồn", level=1)
add_table(
    doc,
    ["Thành phần", "Công nghệ / file mã nguồn"],
    [
        ("Backend", "Java Spring Boot 3.5.14, Spring MVC Controller"),
        ("View", "Thymeleaf template: menu.html, menu-form.html, pos.html, login.html"),
        ("Database", "MySQL, script cafe_management_mysql.sql"),
        ("ORM", "Spring Data JPA, các entity Product, Category, CafeOrder, OrderItem, Payment, Receipt, User"),
        ("Security", "BCryptPasswordEncoder, session login tự xây dựng trong LoginController"),
        ("Build tool", "Maven Wrapper: mvnw.cmd, pom.xml"),
        ("Ngôn ngữ UI", "HTML, CSS, JavaScript thuần"),
    ],
)

doc.add_heading("2.1. Các file chính liên quan đến phạm vi báo cáo", level=2)
add_table(
    doc,
    ["File", "Vai trò"],
    [
        ("Nhom2Application.java", "Điểm khởi động ứng dụng Spring Boot."),
        ("DataLoader.java", "Khởi tạo tài khoản admin/staff, danh mục và sản phẩm mẫu."),
        ("MenuController.java", "Xử lý màn hình quản lý menu, thêm/sửa/xóa sản phẩm, tìm kiếm và lọc theo danh mục."),
        ("PosController.java", "Hiển thị POS và xử lý checkout tạo đơn hàng."),
        ("Product.java", "Entity sản phẩm, liên kết với Category."),
        ("Category.java", "Entity danh mục sản phẩm."),
        ("CafeOrder.java", "Entity đơn hàng."),
        ("OrderItem.java", "Entity chi tiết đơn hàng, có size/ice/sugar/temperature."),
        ("menu.html", "Giao diện quản lý thực đơn cho admin."),
        ("pos.html", "Giao diện bán hàng tại quầy cho staff."),
        ("cafe_management_mysql.sql", "Thiết kế schema MySQL cho dự án."),
    ],
)

doc.add_heading("3. Áp dụng mô hình Incremental", level=1)
doc.add_paragraph(
    "Mô hình Incremental chia hệ thống thành các phiên bản tăng trưởng. Mỗi phiên bản bổ sung một phần chức năng có giá trị sử dụng, "
    "không chờ đến cuối dự án mới có sản phẩm chạy được. Với dự án quản lý quán cà phê, cách chia version phù hợp vì nghiệp vụ có thể tách thành từng phần độc lập: quản lý menu trước, sau đó dùng menu để tạo order tại POS."
)
add_table(
    doc,
    ["Version", "Increment", "Mục tiêu", "Kết quả mong đợi"],
    [
        ("Version 1", "Menu Management Increment", "Thiết lập project và dữ liệu nền; quản lý sản phẩm trong menu.", "Admin đăng nhập, xem danh sách sản phẩm, thêm/sửa/xóa sản phẩm, lọc theo danh mục."),
        ("Version 2", "Core POS Order Increment", "Bổ sung nghiệp vụ bán hàng tại quầy.", "Staff mở POS, chọn món, tăng/giảm số lượng, checkout và lưu đơn hàng."),
        ("Version sau", "Reporting/Payment/Inventory Increment", "Mở rộng báo cáo, thanh toán chi tiết, hóa đơn, tồn kho.", "Hệ thống hoàn thiện hơn cho vận hành thực tế."),
    ],
)

doc.add_heading("3.1. Lý do chọn Incremental", level=2)
add_bullets(
    doc,
    [
        "Version 1 tạo dữ liệu menu là điều kiện bắt buộc để Version 2 có thể bán hàng.",
        "Mỗi version có thể demo độc lập và nhận phản hồi sớm từ giảng viên/người dùng.",
        "Rủi ro được giảm vì nhóm kiểm tra từng phần nhỏ thay vì tích hợp toàn bộ vào cuối kỳ.",
        "Chức năng sau kế thừa dữ liệu và cấu trúc của chức năng trước, đúng tinh thần tăng trưởng dần.",
    ],
)

doc.add_heading("4. Product Backlog", level=1)
add_table(
    doc,
    ["ID", "User Story", "Độ ưu tiên", "Version"],
    [
        ("US01", "Là admin, tôi muốn đăng nhập hệ thống để truy cập khu vực quản trị.", "High", "V1"),
        ("US02", "Là admin, tôi muốn xem danh sách sản phẩm để biết menu hiện có.", "High", "V1"),
        ("US03", "Là admin, tôi muốn thêm sản phẩm mới vào menu.", "High", "V1"),
        ("US04", "Là admin, tôi muốn sửa thông tin sản phẩm khi giá, mô tả hoặc trạng thái thay đổi.", "High", "V1"),
        ("US05", "Là admin, tôi muốn xóa sản phẩm không còn bán.", "Medium", "V1"),
        ("US06", "Là admin, tôi muốn lọc sản phẩm theo danh mục và tìm kiếm theo tên/mô tả.", "Medium", "V1"),
        ("US07", "Là hệ thống, tôi muốn lưu sản phẩm và danh mục vào MySQL để dữ liệu không mất sau khi tắt ứng dụng.", "High", "V1"),
        ("US08", "Là nhân viên, tôi muốn mở giao diện POS để bán hàng tại quầy.", "High", "V2"),
        ("US09", "Là nhân viên, tôi muốn chọn món từ menu để thêm vào đơn hiện tại.", "High", "V2"),
        ("US10", "Là nhân viên, tôi muốn tăng/giảm số lượng món trong giỏ hàng.", "High", "V2"),
        ("US11", "Là nhân viên, tôi muốn hệ thống tính tổng tiền tự động.", "High", "V2"),
        ("US12", "Là nhân viên, tôi muốn checkout tiền mặt để lưu đơn hàng.", "High", "V2"),
        ("US13", "Là hệ thống, tôi muốn lưu chi tiết đơn gồm sản phẩm, số lượng và tùy chỉnh size/đá/đường/nhiệt độ.", "Medium", "V2"),
    ],
)

doc.add_heading("5. Definition of Done chung", level=1)
doc.add_heading("5.1. DoD kỹ thuật", level=2)
add_bullets(
    doc,
    [
        "Code compile được bằng Maven.",
        "Không có lỗi syntax ở controller, model, repository và template chính.",
        "Chức năng có mapping URL rõ ràng trong controller.",
        "Dữ liệu nghiệp vụ được map với entity JPA và lưu được xuống MySQL.",
        "Chức năng mới không làm hỏng chức năng của version trước.",
    ],
)
doc.add_heading("5.2. DoD giao diện", level=2)
add_bullets(
    doc,
    [
        "Trang HTML/Thymeleaf hiển thị đúng dữ liệu từ database.",
        "Form nhập liệu có các trường cần thiết và submit về đúng endpoint.",
        "Nút thêm, sửa, xóa, chọn món và thanh toán hoạt động theo luồng nghiệp vụ.",
        "Có lọc/tìm kiếm cơ bản để người dùng thao tác nhanh.",
    ],
)
doc.add_heading("5.3. DoD dữ liệu", level=2)
add_bullets(
    doc,
    [
        "Có bảng MySQL tương ứng cho danh mục, sản phẩm, đơn hàng và chi tiết đơn hàng.",
        "Các quan hệ chính được thể hiện bằng khóa ngoại hoặc annotation JPA.",
        "Dữ liệu mẫu đủ để demo sau khi chạy ứng dụng.",
    ],
)

doc.add_heading("6. Version 1: Quản lý Menu", level=1)
doc.add_heading("6.1. Version Goal", level=2)
doc.add_paragraph(
    "Thiết lập nền tảng Spring Boot cho dự án và xây dựng module quản lý menu để admin có thể quản lý sản phẩm bán tại quán. "
    "Đây là increment đầu tiên vì dữ liệu sản phẩm/danh mục là đầu vào trực tiếp cho POS ở Version 2."
)
doc.add_heading("6.2. Version Backlog", level=2)
add_table(
    doc,
    ["User Story", "Task triển khai", "File/Thành phần"],
    [
        ("US01", "Tạo login và lưu session userId, role.", "LoginController.java, login.html, User.java"),
        ("US02", "Hiển thị danh sách sản phẩm, danh mục, trạng thái.", "MenuController.menuPage(), menu.html"),
        ("US03", "Tạo form/modal thêm sản phẩm.", "GET/POST /menu/add, menu.html, menu-form.html"),
        ("US04", "Tạo form/modal sửa sản phẩm.", "GET/POST /menu/edit/{id}, menu.html, menu-form.html"),
        ("US05", "Xóa sản phẩm theo id.", "GET /menu/delete/{id}"),
        ("US06", "Tìm kiếm sản phẩm theo tên/mô tả và lọc theo danh mục.", "RequestParam category, search trong MenuController"),
        ("US07", "Tạo entity/repository/schema cho Category và Product.", "Category.java, Product.java, CategoryRepository.java, ProductRepository.java, cafe_management_mysql.sql"),
    ],
)
doc.add_heading("6.3. Thiết kế dữ liệu của Version 1", level=2)
add_code(
    doc,
    """CREATE TABLE categories (
    category_id INT AUTO_INCREMENT PRIMARY KEY,
    category_name VARCHAR(100) NOT NULL,
    description TEXT,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
);

CREATE TABLE products (
    product_id INT AUTO_INCREMENT PRIMARY KEY,
    category_id INT,
    product_name VARCHAR(100),
    description TEXT,
    base_price DECIMAL(10, 2) NOT NULL,
    image_url VARCHAR(255),
    status ENUM('ACTIVE', 'OUT_OF_STOCK', 'INACTIVE') DEFAULT 'ACTIVE',
    FOREIGN KEY (category_id) REFERENCES categories(category_id)
);""",
)
doc.add_heading("6.4. Chức năng giao diện Version 1", level=2)
add_bullets(
    doc,
    [
        "Sidebar quản trị gồm Dashboard, Thực đơn, Lịch sử đơn hàng, Báo cáo và Đăng xuất.",
        "Màn hình menu hiển thị sản phẩm dạng lưới, gồm ảnh, tên món, giá, mô tả, danh mục và trạng thái.",
        "Admin có thể tìm kiếm món theo tên hoặc mô tả.",
        "Admin có thể lọc món theo danh mục như Coffee, Tea, Juice, Cake, Food, Other.",
        "Admin có thể thêm sản phẩm mới với tên, giá, danh mục, mô tả, URL ảnh và trạng thái.",
        "Admin có thể sửa sản phẩm bằng modal hoặc trang form.",
        "Admin có thể xóa sản phẩm sau khi xác nhận.",
    ],
)
doc.add_heading("6.5. Increment sau Version 1", level=2)
doc.add_paragraph("Increment 1: Menu Management Module chạy được.")
add_bullets(
    doc,
    [
        "Project Spring Boot đã có cấu trúc MVC cơ bản.",
        "MySQL có bảng categories và products.",
        "Ứng dụng có dữ liệu mẫu được seed bằng DataLoader.",
        "Admin có thể thêm, xem, sửa, xóa sản phẩm.",
        "Danh mục được dùng để phân loại, lọc và chọn khi tạo/sửa sản phẩm.",
        "Sản phẩm có trạng thái ACTIVE, OUT_OF_STOCK, INACTIVE để phục vụ bán hàng.",
    ],
)
doc.add_heading("6.6. DoD Version 1", level=2)
add_table(
    doc,
    ["Tiêu chí", "Kết quả theo code hiện tại"],
    [
        ("Project Spring Boot compile được", "Đạt - đã kiểm tra bằng mvnw.cmd -q -DskipTests compile."),
        ("Có bảng products và categories trong MySQL script", "Đạt."),
        ("Có entity/repository Product và Category", "Đạt."),
        ("Hiển thị danh sách sản phẩm", "Đạt - GET /menu."),
        ("Tìm kiếm và lọc sản phẩm theo danh mục", "Đạt - xử lý trong MenuController và menu.html."),
        ("Thêm sản phẩm", "Đạt - POST /menu/add."),
        ("Sửa sản phẩm", "Đạt - POST /menu/edit/{id}."),
        ("Xóa sản phẩm", "Đạt - GET /menu/delete/{id}."),
        ("CRUD danh mục riêng", "Một phần - code có bảng/entity/repository và dữ liệu seed danh mục; chưa có màn hình/controller riêng để thêm/sửa/xóa danh mục."),
    ],
)
doc.add_heading("6.7. Review và Retrospective Version 1", level=2)
doc.add_paragraph("Review: Version 1 có thể demo bằng cách đăng nhập admin, mở trang /menu, thêm món mới, sửa giá/mô tả/trạng thái và xóa sản phẩm.")
doc.add_paragraph("Điều làm tốt:")
add_bullets(
    doc,
    [
        "Cấu trúc MVC rõ ràng, controller tách riêng cho menu.",
        "Dữ liệu sản phẩm gắn với danh mục, sẵn sàng dùng lại ở POS.",
        "Giao diện menu trực quan hơn dạng bảng thuần vì dùng card sản phẩm.",
    ],
)
doc.add_paragraph("Vấn đề còn lại và cải tiến:")
add_bullets(
    doc,
    [
        "Nên bổ sung CategoryController và giao diện CRUD danh mục nếu yêu cầu chấm điểm bắt buộc có thêm/sửa/xóa danh mục độc lập.",
        "Nên chuẩn hóa hiển thị tiền tệ VND vì DataLoader đang dùng một số giá dạng thập phân.",
        "Nên đổi thao tác xóa từ GET sang POST/DELETE để phù hợp REST và an toàn hơn.",
    ],
)

doc.add_heading("7. Version 2: Tạo Order cốt lõi", level=1)
doc.add_heading("7.1. Version Goal", level=2)
doc.add_paragraph(
    "Bổ sung giao diện POS để nhân viên bán hàng tại quầy có thể chọn món từ menu đã có ở Version 1, tạo giỏ hàng, tính tổng tiền và checkout. "
    "Sau Version 2, quán có thể dùng thử hệ thống ở mức cơ bản để tính tiền và lưu đơn hàng."
)
doc.add_heading("7.2. Version Backlog", level=2)
add_table(
    doc,
    ["User Story", "Task triển khai", "File/Thành phần"],
    [
        ("US08", "Tạo trang POS cho nhân viên sau đăng nhập.", "GET /pos, pos.html"),
        ("US09", "Load danh sách sản phẩm và danh mục từ database.", "ProductRepository.findAll(), CategoryRepository.findAll()"),
        ("US09", "Chọn món và thêm vào giỏ hàng bằng JavaScript.", "Nút Thêm món trong pos.html"),
        ("US10", "Tăng/giảm số lượng món trong đơn hiện tại.", "qty-controls trong pos.html"),
        ("US11", "Tự động tính tổng tiền theo giá x số lượng.", "renderOrder(), formatMoney() trong pos.html"),
        ("US12", "Checkout và gửi cartJson về backend.", "POST /pos/checkout"),
        ("US13", "Lưu order và order_items xuống database.", "CafeOrder.java, OrderItem.java, PosController.checkout()"),
    ],
)
doc.add_heading("7.3. Thiết kế dữ liệu của Version 2", level=2)
add_code(
    doc,
    """CREATE TABLE orders (
    order_id INT AUTO_INCREMENT PRIMARY KEY,
    staff_id INT,
    total_amount DECIMAL(10, 2) NOT NULL,
    order_status ENUM('Preparing', 'Ready', 'Served', 'Cancelled') DEFAULT 'Preparing',
    order_date TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    FOREIGN KEY (staff_id) REFERENCES users(user_id) ON DELETE SET NULL
);

CREATE TABLE order_items (
    order_item_id INT AUTO_INCREMENT PRIMARY KEY,
    order_id INT NOT NULL,
    product_id INT NOT NULL,
    quantity INT NOT NULL DEFAULT 1,
    size ENUM('S', 'M', 'L') DEFAULT 'M',
    ice_level ENUM('0%', '30%', '50%', '70%', '100%') DEFAULT '100%',
    sugar_level ENUM('0%', '30%', '50%', '70%', '100%') DEFAULT '100%',
    temperature ENUM('ICE', 'HOT') DEFAULT 'ICE',
    FOREIGN KEY (order_id) REFERENCES orders(order_id),
    FOREIGN KEY (product_id) REFERENCES products(product_id)
);""",
)
doc.add_heading("7.4. Luồng xử lý checkout", level=2)
add_numbers(
    doc,
    [
        "Nhân viên đăng nhập tài khoản staff và được chuyển đến /pos.",
        "Backend nạp products và categories vào model để render giao diện POS.",
        "Nhân viên tìm kiếm/lọc danh mục và bấm Thêm món.",
        "JavaScript thêm món vào mảng cart, cho phép tăng/giảm số lượng và tính tổng tiền.",
        "Khi bấm Thanh toán, form gửi cartJson lên POST /pos/checkout.",
        "PosController đọc cartJson bằng ObjectMapper, tạo CafeOrder mới và gán staff theo session userId.",
        "Với mỗi item trong cart, hệ thống tìm Product, tạo OrderItem, set quantity, size, iceLevel, sugarLevel, temperature.",
        "Tổng tiền được tính bằng basePrice x quantity và lưu vào CafeOrder.totalAmount.",
        "Order được lưu bằng OrderRepository; các OrderItem được lưu cascade theo quan hệ OneToMany.",
        "Sau khi lưu thành công, người dùng được redirect về /pos để tiếp tục bán hàng.",
    ],
)
doc.add_heading("7.5. Increment sau Version 2", level=2)
doc.add_paragraph("Increment 2: Core POS Order Module chạy được.")
add_bullets(
    doc,
    [
        "Nhân viên có giao diện bán hàng riêng.",
        "POS tái sử dụng danh mục và sản phẩm từ Version 1.",
        "Có tìm kiếm món và lọc theo danh mục.",
        "Có giỏ hàng hiện tại, tăng/giảm số lượng và xóa giỏ hàng.",
        "Tổng tiền được tính tự động theo giá sản phẩm.",
        "Checkout lưu đơn hàng và chi tiết đơn hàng vào database.",
        "Các trường size, iceLevel, sugarLevel, temperature đã có ở entity/backend để chuẩn bị cho tùy chỉnh món.",
    ],
)
doc.add_heading("7.6. DoD Version 2", level=2)
add_table(
    doc,
    ["Tiêu chí", "Kết quả theo code hiện tại"],
    [
        ("Trang POS hiển thị danh sách món", "Đạt - GET /pos load products và categories."),
        ("Lọc danh mục và tìm kiếm món tại POS", "Đạt - xử lý bằng JavaScript trên pos.html."),
        ("Chọn món vào giỏ hàng", "Đạt."),
        ("Tăng/giảm số lượng món", "Đạt."),
        ("Tính tổng tiền tự động", "Đạt."),
        ("Checkout tạo order", "Đạt - POST /pos/checkout tạo CafeOrder."),
        ("Lưu chi tiết order_items", "Đạt - tạo OrderItem và lưu cascade theo CafeOrder."),
        ("Tùy chỉnh đá/đường/size/nhiệt độ", "Một phần - backend và database có field; giao diện hiện tại gửi mặc định M, 100% đá, 100% đường, ICE, chưa có control chọn trực tiếp."),
        ("Thanh toán tiền mặt", "Một phần - POS có nút Thanh toán và lưu đơn; Payment entity/schema có CASH nhưng checkout hiện chưa tạo bản ghi Payment riêng."),
    ],
)
doc.add_heading("7.7. Review và Retrospective Version 2", level=2)
doc.add_paragraph("Review: Version 2 có thể demo bằng cách đăng nhập staff, mở /pos, chọn nhiều món, thay đổi số lượng, quan sát tổng tiền và bấm Thanh toán để lưu order.")
doc.add_paragraph("Điều làm tốt:")
add_bullets(
    doc,
    [
        "POS tận dụng lại dữ liệu menu từ Version 1, đúng tính chất incremental.",
        "Giao diện chia rõ vùng danh sách món và vùng đơn hiện tại, phù hợp thao tác bán hàng tại quầy.",
        "Backend checkout dùng transaction, hạn chế lưu dở khi có lỗi trong quá trình tạo đơn.",
    ],
)
doc.add_paragraph("Vấn đề còn lại và cải tiến:")
add_bullets(
    doc,
    [
        "Bổ sung control chọn size, mức đá, mức đường và nóng/lạnh trên từng item trong giỏ hàng.",
        "Khi checkout tiền mặt nên tạo thêm Payment với paymentMethod = CASH, paymentStatus = COMPLETED và paymentDate = thời điểm thanh toán.",
        "Nên chặn sản phẩm OUT_OF_STOCK/INACTIVE ở backend để tránh người dùng gửi request thủ công.",
        "Nên hiển thị thông báo thanh toán thành công thay vì chỉ redirect về POS.",
    ],
)

doc.add_heading("8. Thiết kế cơ sở dữ liệu và kiến trúc xử lý", level=1)
doc.add_heading("8.1. Các bảng chính", level=2)
add_table(
    doc,
    ["Bảng", "Mục đích", "Quan hệ chính"],
    [
        ("users", "Lưu tài khoản admin/staff/customer.", "Một staff có thể tạo nhiều orders."),
        ("categories", "Lưu danh mục món.", "Một category có nhiều products."),
        ("products", "Lưu món trong menu.", "Một product thuộc một category; xuất hiện trong nhiều order_items."),
        ("orders", "Lưu đơn hàng.", "Một order có nhiều order_items; gắn với staff."),
        ("order_items", "Lưu chi tiết từng món trong đơn.", "Gắn với orders và products."),
        ("payments", "Chuẩn bị lưu thanh toán CASH/VNPAY.", "Gắn với orders."),
        ("receipts", "Chuẩn bị lưu nội dung hóa đơn.", "Gắn với orders."),
    ],
)
doc.add_heading("8.2. Quan hệ dữ liệu", level=2)
add_code(
    doc,
    """categories 1 ---- n products
users      1 ---- n orders (staff_id)
orders     1 ---- n order_items
products   1 ---- n order_items
orders     1 ---- n payments
orders     1 ---- n receipts""",
)
doc.add_heading("8.3. Kiến trúc xử lý theo MVC", level=2)
add_table(
    doc,
    ["Lớp", "Thành phần trong code", "Nhiệm vụ"],
    [
        ("View", "Thymeleaf templates: menu.html, pos.html", "Hiển thị UI, nhận thao tác người dùng, gửi form/request."),
        ("Controller", "MenuController, PosController, LoginController", "Điều phối request, kiểm tra session/role, gọi repository và trả view/redirect."),
        ("Model", "Product, Category, CafeOrder, OrderItem, User", "Biểu diễn dữ liệu nghiệp vụ và mapping với bảng MySQL."),
        ("Repository", "JpaRepository interfaces", "Thao tác CRUD với database thông qua Spring Data JPA."),
        ("Database", "MySQL schema", "Lưu dữ liệu thực đơn, người dùng, đơn hàng và thanh toán."),
    ],
)

doc.add_heading("9. Kiểm thử và đánh giá hiện trạng", level=1)
doc.add_heading("9.1. Kiểm thử đã thực hiện", level=2)
add_table(
    doc,
    ["Hạng mục", "Cách kiểm thử", "Kết quả"],
    [
        ("Compile project", "Chạy mvnw.cmd -q -DskipTests compile trong thư mục nhom2.", "Đạt."),
        ("Login role admin/staff", "Kiểm tra LoginController và DataLoader tạo tài khoản admin/staff.", "Có luồng xử lý."),
        ("Menu CRUD sản phẩm", "Rà endpoint /menu, /menu/add, /menu/edit/{id}, /menu/delete/{id}.", "Có code xử lý."),
        ("POS checkout", "Rà endpoint /pos và /pos/checkout.", "Có code tạo order và order item."),
        ("Database mapping", "Đối chiếu entity với cafe_management_mysql.sql.", "Phù hợp ở các bảng chính; có một số field schema chưa được dùng đầy đủ trong code."),
    ],
)
doc.add_heading("9.2. Kịch bản demo đề xuất", level=2)
add_numbers(
    doc,
    [
        "Chạy MySQL và import file cafe_management_mysql.sql.",
        "Chạy ứng dụng Spring Boot bằng Maven.",
        "Đăng nhập admin/admin123, mở /menu.",
        "Thêm một sản phẩm mới, chọn danh mục và trạng thái ACTIVE.",
        "Sửa giá hoặc mô tả sản phẩm vừa thêm.",
        "Tìm kiếm sản phẩm và lọc theo danh mục.",
        "Đăng xuất, đăng nhập staff/staff123.",
        "Mở /pos, chọn sản phẩm, tăng/giảm số lượng.",
        "Bấm Thanh toán và kiểm tra bảng orders/order_items trong MySQL.",
    ],
)
doc.add_heading("9.3. Các điểm cần hoàn thiện để khớp 100% yêu cầu được giao", level=2)
add_table(
    doc,
    ["Yêu cầu", "Hiện trạng code", "Đề xuất hoàn thiện"],
    [
        ("Thêm/Sửa/Xóa danh mục", "Có entity, repository, seed data và dùng danh mục trong UI; chưa có CRUD danh mục riêng.", "Thêm CategoryController, category.html hoặc modal quản lý danh mục trong menu."),
        ("Chọn tùy chỉnh đá/đường", "OrderItem có field iceLevel/sugarLevel/temperature; POS đang gửi giá trị mặc định.", "Thêm select/radio cho size, ice, sugar, temperature trong cart item."),
        ("Thanh toán tiền mặt", "Checkout lưu order; Payment entity/schema có CASH nhưng chưa được tạo khi checkout.", "Inject PaymentRepository và save Payment sau khi order được lưu."),
        ("Thông báo sau checkout", "Redirect về /pos.", "Dùng redirect attribute hoặc toast để báo thanh toán thành công."),
    ],
)

doc.add_heading("10. Tổng hợp Increment", level=1)
add_table(
    doc,
    ["Version", "Increment", "Mô tả", "Khả năng sử dụng"],
    [
        ("Version 1", "Menu Management", "Admin quản lý sản phẩm, lọc/tìm kiếm theo danh mục, dữ liệu lưu MySQL.", "Có thể dùng để chuẩn bị và cập nhật thực đơn."),
        ("Version 2", "Core POS Order", "Staff chọn món, lập giỏ hàng, tính tổng tiền và lưu order.", "Có thể dùng thử để bán hàng tại quầy ở mức cơ bản."),
    ],
)

doc.add_heading("11. Kết luận", level=1)
doc.add_paragraph(
    "Dự án quản lý quán cà phê phù hợp với mô hình Incremental vì mỗi version đều tạo thêm một phần giá trị sử dụng rõ ràng. "
    "Version 1 tạo nền tảng dữ liệu menu và chức năng quản trị sản phẩm. Version 2 kế thừa menu để xây dựng POS và tạo đơn hàng cốt lõi. "
    "Hai increment này tạo thành nền móng để nhóm tiếp tục mở rộng sang thanh toán chi tiết, hóa đơn, lịch sử đơn hàng, báo cáo doanh thu và các chức năng quản trị nâng cao."
)
doc.add_paragraph(
    "Theo mã nguồn hiện tại, các phần sản phẩm và tạo order đã có cấu trúc rõ ràng và compile thành công. "
    "Một số yêu cầu như CRUD danh mục riêng, chọn tùy chỉnh đá/đường trên UI và lưu bản ghi Payment tiền mặt cần được bổ sung để hệ thống khớp hoàn toàn với mô tả ban đầu."
)

doc.add_heading("Phụ lục: Tài khoản demo từ DataLoader", level=1)
add_table(
    doc,
    ["Vai trò", "Username", "Password", "Điều hướng sau đăng nhập"],
    [
        ("Admin", "admin", "admin123", "/admin"),
        ("Staff", "staff", "staff123", "/pos"),
    ],
)

p = doc.add_paragraph("Báo cáo được lập dựa trên mã nguồn trong thư mục: E:\\QLDAPM\\Management-Application-Project-Group-2\\nhom2")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    set_run_font(run, size=11, italic=True)

OUT.parent.mkdir(exist_ok=True)
doc.save(OUT)
print(OUT)
